from docx import Document
from docx.shared import Inches

students=[{'std_name':'vivek','roll_no':1,'class':'12','div':'A',
           'ENG':55,'MTH':68,'PHY':56,'CHE':60,'ELE':79,'TMO':318,'TMK':500,'sub_TM':'100'},
          {'std_name':'vivek','roll_no':2,'class':'12','div':'B',
           'ENG':56,'MTH':38,'PHY':57,'CHE':61,'ELE':80,'TMO':323,'TMK':500,'sub_TM':'100'}]

def results(a):
    if students [a]['ENG']< 40:
        result='Fail'
    elif students [a]['MTH'] < 40:
        result='Fail'
    elif students [a]['PHY'] < 40:
        result='Fail'
    elif students [a]['CHE'] < 40:
        result='Fail'
    elif students [a]['ELE'] < 40:
        result='Fail'
    else:
        result='pass'
    return result

def table_m(a,b,c):
    row_cells = table.add_row().cells
    row_cells[0].text = a
    row_cells[1].text = students[b]['sub_TM']
    row_cells[2].text =str(c)

def percent(a):
    i=int(a)
    o=students[i]['TMO']
    t=students[i]['TMK']
    per=(o/t)*100
    return per
for i in range(0, len(students)):
    percnt= percent(a=i)
    reslt = results(a=i)
    std_nam=students [i]['std_name']
    std_nam= Document()
    std_nam.add_heading('Result', 0)
    std_nam.add_heading('Class %s Result'%students [i]['class'])
    p = std_nam.add_paragraph('Roll Number:%d\t\t\t'%students [i]['roll_no'])
    p.add_run('Student Name :%s\n\n'%students [i]['std_name'])
    p.add_run('class :%s\t\t\t'%students [i]['class'])
    p.add_run('Division :%s\n'%students[i]['div'])

    table = std_nam.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Subjets'
    hdr_cells[1].text = 'Max Marks'
    hdr_cells[2].text = 'Obtained Marks'

    table_m(a='English', b=i , c=students[i]['ENG'])
    table_m(a='Mathmatics', b=i , c=students[i]['MTH'])
    table_m(a='Physics', b= i, c=students[i]['PHY'])
    table_m(a='Chemistry', b=i , c=students[i]['CHE'])
    table_m(a='Electronices', b=i , c=students[i]['ELE'])

    row_cells = table.add_row().cells
    row_cells[0].text ='Total Marks'
    row_cells[1].text =str(students [i]['TMK'])
    row_cells[2].text =str(students [i]['TMO'])

    p = std_nam.add_paragraph('\n\nResult : %s \t\t\t'%reslt)
    p.add_run('Percentage :%d'%percnt)

    std_nam.add_page_break()
    std_nam.save(''+students[i]['std_name']+'_'+str(students[i]['roll_no'])+'.docx')
    


